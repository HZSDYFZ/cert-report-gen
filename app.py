import sys, io, re, zipfile
from datetime import datetime
import pandas as pd
import streamlit as st
from docx import Document

CHK_EMPTY = chr(0x25A1)
CHK_FILLED = chr(0x25A0)

def format_date(val):
    if val is None: return ''
    if isinstance(val, datetime): return val.strftime('%Y-%m-%d')
    s = str(val).strip()
    m = re.search(r'(\d{4})[年\-/](\d{1,2})[月\-/](\d{1,2})', s)
    if m: return m.group(1)+'-'+m.group(2).zfill(2)+'-'+m.group(3).zfill(2)
    m = re.search(r'(\d{1,2})/(\d{1,2})/(\d{4})', s)
    if m: return m.group(3)+'-'+m.group(1).zfill(2)+'-'+m.group(2).zfill(2)
    return s[:10] if len(s) >= 10 else s

def get_conclusion_idx(atype):
    atype = str(atype).strip() if atype else ''
    is_surv = '监' in atype and '再认证' not in atype and '二阶段' not in atype and '一阶段' not in atype
    if '一阶段' in atype or '二阶段' in atype or '再认证' in atype: return 0
    elif '转移' in atype: return 2
    elif is_surv:
        if '换发' in atype: return 5
        else: return 4
    else: return 0

def is_audit_surv(atype):
    return '监' in atype and '再认证' not in atype and '二阶段' not in atype and '一阶段' not in atype

def set_fcb(doc_xml, pos, checked):
    cs = max(0, pos - 300)
    chunk = doc_xml[cs:pos + 100]
    if checked:
        if 'w:checked w:val=\"0\"' in chunk:
            return doc_xml[:cs] + chunk.replace('w:checked w:val=\"0\"/>', 'w:checked/>') + doc_xml[pos + 100:]
        if 'w:checked' not in chunk:
            return doc_xml[:cs] + chunk.replace('FORMCHECKBOX', 'w:checked/>FORMCHECKBOX') + doc_xml[pos + 100:]
    else:
        if 'w:checked/>' in chunk:
            return doc_xml[:cs] + chunk.replace('w:checked/>', 'w:checked w:val=\"0\"/>') + doc_xml[pos + 100:]
    return doc_xml

def fill_cert_standard(cell_text, atype):
    at = atype.strip()
    result = cell_text
    if 'IATF16949' in cell_text and 'IATF' in at:
        result = result.replace(CHK_EMPTY + ' IATF16949', CHK_FILLED + ' IATF16949')
    if 'ISO9001' in cell_text and '9001' in at:
        result = result.replace(CHK_EMPTY + ' ISO9001', CHK_FILLED + ' ISO9001')
    if 'ISO14001' in cell_text and ('EMS' in at or '14001' in at):
        result = result.replace(CHK_EMPTY + ' ISO14001', CHK_FILLED + ' ISO14001')
    if 'ISO45001' in cell_text:
        if 'OHS' in at or '45001' in at:
            result = result.replace(CHK_EMPTY + 'ISO 45001', CHK_FILLED + 'ISO 45001')
            if result == cell_text:
                result = result.replace(CHK_EMPTY + 'ISO45001', CHK_FILLED + 'ISO45001')
    return result

def fill_audit_type(cell_text, atype):
    at = atype.strip()
    is_surv = is_audit_surv(at)
    result = cell_text
    if '一阶段' in at or '二阶段' in at or '再认证' in at:
        result = result.replace(CHK_EMPTY + '初审', CHK_FILLED + '初审')
    if is_surv:
        result = result.replace(CHK_EMPTY + '监审', CHK_FILLED + '监审')
    if '再认证' in at or '转移' in at:
        result = result.replace(CHK_EMPTY + '再认证/转移', CHK_FILLED + '再认证/转移')
    if '特殊' in at:
        result = result.replace(CHK_EMPTY + '特殊审核', CHK_FILLED + '特殊审核')
    return result

def fill_report(doc, fields):
    company = fields.get('company', '')
    taskNo = fields.get('taskNo', '')
    leader = fields.get('leader', '')
    auditType = fields.get('auditType', '')
    address = fields.get('address', '')
    scope = fields.get('scope', '')
    filled = {'company': False, 'taskNo': False, 'leader': False, 'address': False, 'scope': False}
    for para in doc.paragraphs:
        runs = list(para.runs)
        full_text = ''.join(r.text or '' for r in runs)
        if company and not filled['company'] and '公司名称' in full_text:
            for i, run in enumerate(runs):
                if run.text and '公司名称' in run.text:
                    if i + 1 < len(runs): runs[i + 1].text = (runs[i + 1].text or '') + company
                    else: run.text = run.text + company
                    filled['company'] = True; break
        if taskNo and not filled['taskNo'] and ('任务号' in full_text or '任务编号' in full_text):
            for i, run in enumerate(runs):
                if run.text and ('任务号' in run.text or '任务编号' in run.text):
                    if i + 1 < len(runs): runs[i + 1].text = (runs[i + 1].text or '') + taskNo
                    else: run.text = run.text + taskNo
                    filled['taskNo'] = True; break
        if leader and not filled['leader'] and '审核组长' in full_text:
            for i, run in enumerate(runs):
                if run.text and '审核组长' in run.text:
                    if i + 1 < len(runs): runs[i + 1].text = (runs[i + 1].text or '') + leader
                    else: run.text = run.text + leader
                    filled['leader'] = True; break
        if address and not filled['address'] and '审核地址' in full_text:
            for i, run in enumerate(runs):
                if run.text and '审核地址' in run.text:
                    if i + 1 < len(runs): runs[i + 1].text = (runs[i + 1].text or '') + address
                    else: run.text = run.text + address
                    filled['address'] = True; break
        if scope and not filled['scope'] and '认证范围' in full_text:
            for i, run in enumerate(runs):
                if run.text and '认证范围' in run.text:
                    if i + 1 < len(runs): runs[i + 1].text = (runs[i + 1].text or '') + scope
                    else: run.text = run.text + scope
                    filled['scope'] = True; break
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            cells = row.cells
            if ri == 0:
                if company and not filled['company']:
                    for ci in range(min(3, len(cells))):
                        for para in cells[ci].paragraphs:
                            runs = list(para.runs)
                            for i, run in enumerate(runs):
                                if run.text and '公司名称' in run.text:
                                    if i + 1 < len(runs): runs[i + 1].text = (runs[i + 1].text or '') + company
                                    else: run.text = run.text + company
                                    filled['company'] = True; break
                            if filled['company']: break
                        if filled['company']: break
                if taskNo and not filled['taskNo'] and len(cells) > 3:
                    for para in cells[3].paragraphs:
                        runs = list(para.runs)
                        for i, run in enumerate(runs):
                            if run.text and ('任务号' in run.text or '任务编号' in run.text):
                                if i + 1 < len(runs): runs[i + 1].text = (runs[i + 1].text or '') + taskNo
                                else: run.text = run.text + taskNo
                                filled['taskNo'] = True; break
                        if filled['taskNo']: break
            if ri == 1 and leader and not filled['leader'] and len(cells) > 2:
                cell = cells[2]
                has_content = any(r.text and r.text.strip() for para in cell.paragraphs for r in para.runs)
                if not has_content:
                    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    para.add_run(leader); filled['leader'] = True
            if ri == 2 and auditType:
                at = auditType.strip()
                for ci in range(2, min(5, len(cells))):
                    for para in cells[ci].paragraphs:
                        new_text = fill_cert_standard(para.text, at)
                        if new_text != para.text:
                            for run in para.runs: run.text = ''
                            if para.runs: para.runs[0].text = new_text
                            else: para.add_run(new_text)
            if ri == 3 and auditType:
                at = auditType.strip()
                for ci in range(2, min(5, len(cells))):
                    for para in cells[ci].paragraphs:
                        new_text = fill_audit_type(para.text, at)
                        if new_text != para.text:
                            for run in para.runs: run.text = ''
                            if para.runs: para.runs[0].text = new_text
                            else: para.add_run(new_text)
            if ri == 4 and address and not filled['address']:
                for ci in range(2, min(5, len(cells))):
                    for para in cells[ci].paragraphs:
                        for run in para.runs:
                            if run.text and '审核地址' in run.text:
                                if run.text.strip() == '审核地址：':
                                    run.text = run.text + address; filled['address'] = True
                                else:
                                    runs = list(para.runs)
                                    for i, r in enumerate(runs):
                                        if r.text and '审核地址' in r.text:
                                            if i + 1 < len(runs): runs[i + 1].text = (runs[i + 1].text or '') + address
                                            else: r.text = r.text + address
                                            filled['address'] = True; break
                                    if filled['address']: break
                                if filled['address']: break
                        if filled['address']: break
            if ri == 5 and scope and not filled['scope']:
                for ci in range(2, min(5, len(cells))):
                    for para in cells[ci].paragraphs:
                        for run in para.runs:
                            if run.text and '认证范围' in run.text:
                                if run.text.strip() == '认证范围：':
                                    run.text = run.text + scope; filled['scope'] = True
                                else:
                                    runs = list(para.runs)
                                    for i, r in enumerate(runs):
                                        if r.text and '认证范围' in r.text:
                                            if i + 1 < len(runs): runs[i + 1].text = (runs[i + 1].text or '') + scope
                                            else: r.text = r.text + scope
                                            filled['scope'] = True; break
                                    if filled['scope']: break
                                if filled['scope']: break
                        if filled['scope']: break
    if auditType:
        con_idx = get_conclusion_idx(auditType)
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        with zipfile.ZipFile(buf, 'r') as z:
            content = {name: z.read(name) for name in z.namelist()}
        doc_xml = content['word/document.xml'].decode('utf-8')
        fcb_positions = []
        for m in re.finditer(r'FORMCHECKBOX', doc_xml):
            pos = m.start()
            cs = max(0, pos - 300)
            chunk = doc_xml[cs:pos + 100]
            if 'w:checked w:val=\"0\"' in chunk: val = '0'
            elif 'w:checked/>' in chunk: val = '1'
            elif 'w:checked' in chunk: val = '1'
            else: val = '?'
            fcb_positions.append((pos, val))
        for idx in range(7):
            abs_idx = idx + 66
            if abs_idx < len(fcb_positions):
                abs_pos, old_val = fcb_positions[abs_idx]
                doc_xml = set_fcb(doc_xml, abs_pos, (idx == con_idx))
        content['word/document.xml'] = doc_xml.encode('utf-8')
        out = io.BytesIO()
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name, data in content.items(): zout.writestr(name, data)
        out.seek(0); doc = Document(out)
    return doc


def _get_col(row, *names, default=""):
    for name in names:
        if name in row.index:
            v = row[name]
            if pd.notna(v):
                s = str(v).strip()
                if s and s.lower() not in ["nan","none","null",""]:
                    return s
    return default

def _is_empty(s):
    s = str(s).strip() if s is not None else ""
    return not s or s.lower() in ["nan","none","null",""]

def _match_row(df, task_no):
    if df.empty or "任务号" not in df.columns:
        return pd.Series()
    if not task_no or task_no not in df["任务号"].values:
        return pd.Series()
    return df[df["任务号"] == task_no].iloc[0]

def _is_empty_series(r):
    return isinstance(r, pd.Series) and r.empty

def parse_and_fix_excel(file_buffer):
    xls = pd.ExcelFile(file_buffer)
    df1 = pd.read_excel(xls, sheet_name="Sheet1")
    df2 = pd.read_excel(xls, sheet_name="Sheet2") if "Sheet2" in xls.sheet_names else pd.DataFrame()
    df3 = pd.read_excel(xls, sheet_name="Sheet3") if "Sheet3" in xls.sheet_names else pd.DataFrame()
    df4 = pd.read_excel(xls, sheet_name="Sheet4") if "Sheet4" in xls.sheet_names else pd.DataFrame()
    if not df3.empty:
        old_cols = list(df3.columns)
        if len(old_cols) >= 1 and old_cols[0] not in ["任务号"]:
            df3 = df3.rename(columns={old_cols[0]: "任务号"})
        if "Observations" not in df3.columns and len(old_cols) >= 2:
            df3 = df3.rename(columns={old_cols[1]: "Sheet3_结论"})
        if "Date" not in df3.columns and len(old_cols) >= 3:
            df3 = df3.rename(columns={old_cols[2]: "Sheet3_日期"})
        df3 = df3.drop_duplicates(subset=["任务号"], keep="first")
    if not df4.empty:
        header_idx = 0
        while header_idx < len(df4) and df4.iloc[header_idx].isna().all():
            header_idx += 1
        if header_idx < len(df4):
            header_row = df4.iloc[header_idx]
            new_cols = {}
            for j, hc in enumerate(header_row):
                if pd.notna(hc):
                    new_cols[j] = str(hc).strip()
            if "File number(s)" in new_cols.values() and "任务号" not in new_cols.values():
                for j, hc in enumerate(header_row):
                    if pd.notna(hc) and str(hc).strip() == "File number(s)":
                        new_cols[j] = "任务号"
                        break
            df4 = df4.rename(columns=new_cols)
            df4 = df4[header_idx + 1:].reset_index(drop=True)
        if "任务号" in df4.columns:
            df4 = df4.drop_duplicates(subset=["任务号"], keep="first")
    if not df2.empty and "任务号" in df2.columns:
        df2 = df2.drop_duplicates(subset=["任务号"], keep="first")
    has_task_in_s1 = "任务号" in df1.columns
    master_list = []
    anomaly_log = []
    for idx, row in df1.iterrows():
        if has_task_in_s1:
            task_no = str(row.get("任务号", "")).strip()
        else:
            task_no = ""
        if _is_empty(task_no) and not df2.empty and "任务号" in df2.columns:
            s1_company = _get_col(row, "客户名称 Client Name", "公司名称", default="")
            if s1_company:
                mask = df2["企业名称"].astype(str).str.contains(s1_company[:8], na=False) | df2["企业中文名字"].astype(str).str.contains(s1_company[:8], na=False)
                match = df2[mask]
                if not match.empty:
                    task_no = str(match.iloc[0]["任务号"]).strip()
        row2 = _match_row(df2, task_no)
        row3 = _match_row(df3, task_no)
        row4 = _match_row(df4, task_no)
        s1_company = _get_col(row, "客户名称 Client Name", "公司名称", default="")
        s2_company = _get_col(row2, "企业中文名字", "企业名称", default="") if not _is_empty_series(row2) else ""
        s4_company = _get_col(row4, "Company name", default="") if not _is_empty_series(row4) else ""
        is_email_polluted = "@" in s1_company if s1_company else False
        if is_email_polluted or not s1_company:
            company_name = s2_company if s2_company else (s4_company if s4_company else "未知企业")
            if is_email_polluted:
                anomaly_log.append({"任务号": task_no, "异常类型": "邮箱污染公司名", "原污染值": s1_company, "修复后值": company_name})
        else:
            company_name = s1_company
        if not _is_empty_series(row2):
            ce = _get_col(row2, "企业英文名字", default="")
            company_en = ce if ce and ce != company_name else ""
        else:
            company_en = s4_company if s4_company != company_name else ""
        if _is_empty(company_en): company_en = ""
        lead = _get_col(row, "审核组长", "审核组 (LA+Co)", default="")
        if not lead and not _is_empty_series(row2):
            lead = _get_col(row2, "组长", default="")
        members = _get_col(row2, "组员", default="") if not _is_empty_series(row2) else ""
        team_str = "%s (成员: %s)" % (lead, members) if (members and members != lead) else lead
        address = _get_col(row, "审核地址", default="")
        is_address_date = re.match(r"^\d{4}[-/]\d{2}[-/]\d{2}", address)
        if is_address_date or not address:
            real_address = _get_col(row2, "审核地址", default="") if not _is_empty_series(row2) else ""
            if is_address_date:
                anomaly_log.append({"任务号": task_no, "异常类型": "地址错入日期", "原污染值": address, "修复后值": real_address})
            address = real_address if real_address else "未填写"
        scope = _get_col(row, "认证范围", default="")
        if not scope and not _is_empty_series(row2):
            scope = _get_col(row2, "审核范围", default="")
        if not scope: scope = ""
        standard = _get_col(row2, "标准", default="ISO/IATF") if not _is_empty_series(row2) else "ISO/IATF"
        if not standard: standard = "ISO/IATF"
        decision = _get_col(row, "认证决定结论", default="")
        if not decision:
            if not _is_empty_series(row3):
                decision = _get_col(row3, "Sheet3_结论", "Observations", default="")
            if not decision and not _is_empty_series(row4):
                decision = _get_col(row4, "Observations", default="")
        date_val = _get_col(row, "日期", "评定通过时间", default="")
        if not date_val:
            if not _is_empty_series(row3):
                date_val = _get_col(row3, "Sheet3_日期", "Date", default="")
            if not date_val and not _is_empty_series(row4):
                date_val = _get_col(row4, "VP pass date", default="")
        date_val = format_date(date_val) if date_val else ""
        master_list.append({
            "序号": _get_col(row, "项目序号 No.", default=str(idx + 1)),
            "合同号": _get_col(row, "合同号 Contract No.", default=""),
            "任务号": task_no,
            "公司中文名": company_name,
            "公司英文名": company_en,
            "审核类型": _get_col(row, "审核类型Audit Type", default=""),
            "认证标准": standard,
            "审核团队": team_str,
            "评定人员": _get_col(row, "评定人员", default=""),
            "审核地址": address,
            "认证范围": scope,
            "认证结论": decision,
            "结论日期": date_val,
        })
    return pd.DataFrame(master_list), pd.DataFrame(anomaly_log)


def extract_form_fields(doc):
    fields = {}
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if ri == 1 and len(cells) > 1 and cells[1]: fields['taskNo'] = cells[1]
            if ri == 2 and len(cells) > 1 and cells[1]: fields['company'] = cells[1]
            if ri == 3 and len(cells) > 1 and cells[1]: fields['leader'] = cells[1]
            if ri == 4 and len(cells) > 1 and cells[1]: fields['auditType'] = cells[1]
            if ri == 5 and len(cells) > 1 and cells[1]: fields['address'] = cells[1]
            if ri == 6 and len(cells) > 1 and cells[1]: fields['scope'] = cells[1]
    for k in ['company','taskNo','leader','auditType','address','scope']:
        fields.setdefault(k, '')
    return fields

def generate_word_zip_batch(df, template_bytes=None):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for idx, row in df.iterrows():
            data_dict = row.to_dict()
            doc = Document(io.BytesIO(template_bytes))
            doc = fill_report(doc, {
                'company': data_dict.get('公司中文名',''),
                'taskNo': data_dict.get('任务号',''),
                'leader': data_dict.get('审核团队','').split('(')[0].strip() if '(' in data_dict.get('审核团队','') else data_dict.get('审核团队',''),
                'auditType': data_dict.get('审核类型',''),
                'address': data_dict.get('审核地址',''),
                'scope': data_dict.get('认证范围',''),
            })
            out = io.BytesIO(); doc.save(out)
            raw_company = str(data_dict.get('公司中文名', f'企业_{idx+1}'))
            company_name = re.sub(r'[\\/*?:\"<>|]', '_', raw_company)
            task_no = re.sub(r'[\\/*?:\"<>|]', '_', str(data_dict.get('任务号', '')))
            filename = f'{company_name}_{task_no}_评定报告.docx' if task_no else f'{company_name}_评定报告.docx'
            zf.writestr(filename, out.getvalue())
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def generate_excel_bytes(df):
    target = io.BytesIO()
    with pd.ExcelWriter(target, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='认证评定解析全量表')
    target.seek(0)
    return target.getvalue()

# ===== Streamlit UI =====
st.set_page_config(page_title='认证评定自动化解析与模版生成系统', layout='wide')
st.title('认证评定自动化解析与模版生成系统')

# Mode selection
mode = st.radio('选择模式', ['单条报告生成 (FORM6101)', '批量生成 (Excel)'], key='mode')

if mode == '单条报告生成 (FORM6101)':
    st.subheader('单条报告生成')
    col1, col2 = st.columns(2)
    with col1:
        form_file = st.file_uploader('上传 FORM6101', type=['docx'], key='form_up')
    with col2:
        tpl_file = st.file_uploader('上传报告模版', type=['docx'], key='tpl_up')
    if form_file and tpl_file:
        if 'form_fields' not in st.session_state or st.session_state.get('form_name') != form_file.name:
            form_bytes = form_file.read()
            st.session_state.form_bytes = form_bytes
            st.session_state.form_name = form_file.name
            st.session_state.form_fields = extract_form_fields(Document(io.BytesIO(form_bytes)))
        if 'tpl_bytes' not in st.session_state or st.session_state.get('tpl_name') != tpl_file.name:
            st.session_state.tpl_bytes = tpl_file.read()
            st.session_state.tpl_name = tpl_file.name
        fields = st.session_state.form_fields
        st.info('提取字段: 公司=' + str(fields.get('company',''))[:25] + ', 任务号=' + str(fields.get('taskNo','')) + ', 组长=' + str(fields.get('leader','')))
        if st.button('生成报告', type='primary', key='gen_single'):
            with st.spinner('生成中...'):
                try:
                    doc = Document(io.BytesIO(st.session_state.tpl_bytes))
                    doc = fill_report(doc, fields)
                    out = io.BytesIO(); doc.save(out); out.seek(0)
                    st.success('报告已生成！')
                    fname = str(fields.get('company', 'report')) + '.docx'
                    st.download_button('下载报告', data=out.getvalue(), file_name=fname,
                                      mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                except Exception as e:
                    st.error('生成失败: ' + str(e))
    elif form_file or tpl_file:
        st.warning('请同时上传 FORM6101 和报告模版')

else:  # 批量生成 (Excel)
    st.subheader('批量生成')
    col1, col2 = st.columns(2)
    with col1:
        excel_file = st.file_uploader('上传 Excel 数据', type=['xlsx'], key='exc_up')
    with col2:
        tpl_file = st.file_uploader('上传报告模版', type=['docx'], key='batch_tpl_up')
    if excel_file and tpl_file:
        state_key = 'batch_expl'
        if state_key not in st.session_state or st.session_state.get('exc_name') != excel_file.name or st.session_state.get('b_tpl_name') != tpl_file.name:
            try:
                with st.spinner('解析 Excel 数据...'):
                    df, anomaly_df = parse_and_fix_excel(excel_file)
                st.session_state.batch_df = df
                st.session_state.batch_anomaly = anomaly_df
                st.session_state.batch_tpl_bytes = tpl_file.read()
                st.session_state.exc_name = excel_file.name
                st.session_state.b_tpl_name = tpl_file.name
                st.session_state.batch_step = 0
                st.session_state.batch_processed = 0
                st.session_state.batch_files = []
                st.session_state.curr_zip = None
            except Exception as e:
                st.error('解析 Excel 失败: ' + str(e))
                st.stop()
        df = st.session_state.batch_df
        anomaly_df = st.session_state.batch_anomaly
        template_bytes = st.session_state.batch_tpl_bytes
        st.success(f'解析完成！共 {len(df)} 条有效记录，{len(anomaly_df)} 条异常记录')
        if not df.empty:
            tab_single, tab_batch = st.tabs(['单条记录生成', '批量生成'])
            with tab_single:
                st.subheader('单条记录生成')
                selected = st.selectbox('选择记录', df['任务号'].tolist())
                if selected:
                    sel_row = df[df['任务号'] == selected].iloc[0]
                    with st.expander('查看数据详情'):
                        st.json(sel_row.to_dict())
                    word_bytes = generate_word_zip_batch(df[df['任务号']==selected], template_bytes)
                    st.download_button('下载 Word 报告', data=word_bytes, file_name=f'{sel_row["公司中文名"]}_评定报告.docx',
                                       mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            with tab_batch:
                st.subheader('批量生成 (每次20个)')
                st.write(f'当前处理进度: **{st.session_state.batch_processed}** / {len(df)} 条')
                if st.button('开始生成 (20个)', key='batch_start'):
                    with st.spinner('生成中...'):
                        step = st.session_state.batch_step
                        start_idx = step * 20
                        end_idx = min(start_idx + 20, len(df))
                        batch_df = df.iloc[start_idx:end_idx]
                        zip_data = generate_word_zip_batch(batch_df, template_bytes)
                        st.session_state.batch_files.extend([zip_data])
                        st.session_state.batch_step += 1
                        st.session_state.batch_processed = end_idx
                        st.session_state.curr_zip = zip_data
                        st.success(f'已生成 {end_idx} / {len(df)} 条报告')
                        st.download_button('下载本批报告 (.zip)', data=zip_data, file_name=f'reports_batch_{step+1}.zip', mime='application/zip')
                        if end_idx >= len(df):
                            st.success('全部报告已生成完毕！')
                            if len(st.session_state.batch_files) > 1:
                                all_zip = io.BytesIO()
                                with zipfile.ZipFile(all_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
                                    for zd in st.session_state.batch_files:
                                        all_zip.write(zd)
                                all_zip.seek(0)
                                st.download_button('下载全部报告 (.zip)', data=all_zip.getvalue(), file_name='all_reports.zip', mime='application/zip')
                if st.button('清除并重试', key='batch_clear'):
                    st.session_state.batch_step = 0
                    st.session_state.batch_processed = 0
                    st.session_state.batch_files = []
                    st.session_state.curr_zip = None
                    st.rerun()
                excel_data = generate_excel_bytes(df)
                st.download_button('导出 Excel 汇总', data=excel_data, file_name='认证评定记录_汇总.xlsx', mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        if not anomaly_df.empty:
            st.warning(f'发现 {len(anomaly_df)} 条数据异常，请查看:')
            st.dataframe(anomaly_df)
    elif excel_file or tpl_file:
        st.warning('请同时上传 Excel 数据和报告模版')
    else:
        st.info('请上传 Excel 数据文件和报告模版')

st.caption('Cert Report Generator v2.10')

